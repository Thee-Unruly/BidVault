"""
extractor.py
────────────
Extracts clean text from any supported document type.
Routes to the right extraction method based on DetectionResult.

Extraction hierarchy:
  Digital PDF  → pdfplumber  (fast, accurate)
  Scanned PDF  → pdf2image + pytesseract  (slower, OCR)
  Mixed PDF    → pdfplumber first, fall back to OCR per page
  Word         → python-docx  (preserves heading structure)
  Text         → direct read
"""

import re
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from .detector import DetectionResult, DocType


@dataclass
class ExtractionResult:
    text: str                           # full cleaned text
    pages: list[str]                    # text per page (for mixed docs)
    metadata: dict = field(default_factory=dict)
    extraction_method: str = ""
    char_count: int = 0
    warnings: list[str] = field(default_factory=list)

    def __post_init__(self):
        self.char_count = len(self.text)


def extract(file_path: str, detection: DetectionResult) -> ExtractionResult:
    """
    Main entry point. Routes to the correct extractor based on doc type.
    """
    path = Path(file_path)

    if detection.doc_type == DocType.DIGITAL_PDF:
        return _extract_digital_pdf(file_path)

    elif detection.doc_type == DocType.SCANNED_PDF:
        return _extract_scanned_pdf(file_path)

    elif detection.doc_type == DocType.MIXED_PDF:
        return _extract_mixed_pdf(file_path)

    elif detection.doc_type == DocType.WORD:
        return _extract_word(file_path)

    elif detection.doc_type == DocType.TEXT:
        return _extract_text(file_path)

    raise ValueError(f"Unknown doc type: {detection.doc_type}")


# ── DIGITAL PDF ───────────────────────────────────────────────────────────────

def _extract_digital_pdf(file_path: str) -> ExtractionResult:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Run: pip install pdfplumber")

    pages = []
    warnings = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""

            # Also extract tables — common in RFP evaluation criteria sections
            tables = page.extract_tables()
            if tables:
                table_text = _tables_to_text(tables)
                text = text + "\n\n" + table_text if text else table_text

            pages.append(clean_text(text))

    full_text = "\n\n--- PAGE BREAK ---\n\n".join(p for p in pages if p.strip())

    return ExtractionResult(
        text               = full_text,
        pages              = pages,
        extraction_method  = "pdfplumber",
        warnings           = warnings,
    )


# ── SCANNED PDF ───────────────────────────────────────────────────────────────

def _extract_scanned_pdf(file_path: str) -> ExtractionResult:
    """
    Convert each page to an image then run Tesseract OCR.
    Requires: pdf2image, pytesseract, Tesseract binary installed.
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        raise ImportError("Run: pip install pdf2image pytesseract && sudo apt install tesseract-ocr poppler-utils")

    # Configure Tesseract path if provided in .env
    tesseract_cmd = os.getenv("TESSERACT_CMD")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    # Configure Poppler path if provided in .env
    poppler_path = os.getenv("POPPLER_PATH")

    pages = []
    warnings = []

    # DPI 300 is the sweet spot: good OCR quality without huge memory use
    images = convert_from_path(file_path, dpi=300, poppler_path=poppler_path)

    for i, image in enumerate(images):
        text = pytesseract.image_to_string(image, lang="eng", config="--psm 3")
        pages.append(clean_text(text))

        if len(pages[-1].strip()) < 50:
            warnings.append(f"Page {i+1}: OCR returned minimal text — scan quality may be poor")

    full_text = "\n\n--- PAGE BREAK ---\n\n".join(p for p in pages if p.strip())

    return ExtractionResult(
        text              = full_text,
        pages             = pages,
        extraction_method = "tesseract_ocr",
        warnings          = warnings,
    )


# ── MIXED PDF ─────────────────────────────────────────────────────────────────

def _extract_mixed_pdf(file_path: str) -> ExtractionResult:
    """
    Handle PDFs where some pages are digital, some are scanned.
    Try pdfplumber per page first; if text is too short, fall back to OCR.
    """
    try:
        import pdfplumber
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        raise ImportError("Run: pip install pdfplumber pdf2image pytesseract")

    # Configure Tesseract path if provided in .env
    tesseract_cmd = os.getenv("TESSERACT_CMD")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    poppler_path = os.getenv("POPPLER_PATH")

    pages    = []
    warnings = []
    methods  = []

    # Pre-render all pages as images (needed for OCR fallback)
    images = convert_from_path(file_path, dpi=300, poppler_path=poppler_path)

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            clean = clean_text(text)

            if len(clean.replace(" ", "")) > 30:
                # Digital page — pdfplumber worked
                pages.append(clean)
                methods.append("pdfplumber")
            else:
                # Scanned page — use OCR on the pre-rendered image
                if i < len(images):
                    ocr_text = pytesseract.image_to_string(images[i], lang="eng", config="--psm 3")
                    pages.append(clean_text(ocr_text))
                    methods.append("ocr")
                else:
                    pages.append("")
                    warnings.append(f"Page {i+1}: could not render for OCR")

    ocr_count = methods.count("ocr")
    if ocr_count:
        warnings.append(f"{ocr_count} pages required OCR")

    full_text = "\n\n--- PAGE BREAK ---\n\n".join(p for p in pages if p.strip())

    return ExtractionResult(
        text              = full_text,
        pages             = pages,
        extraction_method = f"mixed (pdfplumber + ocr, {ocr_count} OCR pages)",
        warnings          = warnings,
    )


# ── WORD ──────────────────────────────────────────────────────────────────────

def _extract_word(file_path: str) -> ExtractionResult:
    """
    Extract from .docx preserving heading structure.
    Headings are marked with [H1], [H2], [H3] tags — used later by the
    structure-aware chunker to split at logical section boundaries.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Run: pip install python-docx")

    doc      = Document(file_path)
    sections = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name.lower()

        if "heading 1" in style:
            sections.append(f"\n[H1] {text}\n")
        elif "heading 2" in style:
            sections.append(f"\n[H2] {text}\n")
        elif "heading 3" in style:
            sections.append(f"\n[H3] {text}\n")
        else:
            sections.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cell_texts:
                rows.append(" | ".join(cell_texts))
        if rows:
            sections.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]\n")

    full_text = "\n".join(sections)

    return ExtractionResult(
        text              = clean_text(full_text),
        pages             = [full_text],   # Word docs don't have natural page breaks
        extraction_method = "python-docx",
    )


# ── PLAIN TEXT ────────────────────────────────────────────────────────────────

def _extract_text(file_path: str) -> ExtractionResult:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return ExtractionResult(
        text              = clean_text(text),
        pages             = [text],
        extraction_method = "plain_text",
    )


# ── UTILITIES ─────────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Normalise extracted text:
    - Collapse excessive whitespace
    - Remove null bytes and control characters
    - Normalise line endings
    - Remove page headers/footers patterns common in Kenyan government docs
    """
    if not text:
        return ""

    # Remove null bytes
    text = text.replace("\x00", "")

    # Normalise line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove repeated special characters (common OCR artifacts)
    text = re.sub(r"[_\-=]{5,}", " ", text)

    # Collapse 3+ consecutive blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Collapse multiple spaces (but preserve intentional indentation)
    text = re.sub(r"[ \t]{3,}", "  ", text)

    # Remove common Kenyan government doc header/footer noise
    # e.g. "Page 1 of 12", "CONFIDENTIAL", repeated document titles
    text = re.sub(r"Page \d+ of \d+", "", text, flags=re.IGNORECASE)

    return text.strip()


def _tables_to_text(tables: list) -> str:
    """Convert pdfplumber table output to readable text."""
    output = []
    for table in tables:
        rows = []
        for row in table:
            cells = [str(cell or "").strip() for cell in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            output.append("[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
    return "\n\n".join(output)
